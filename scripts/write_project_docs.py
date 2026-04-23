from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


README_TEXT = """
# K线形态匹配工具

一个面向金融课程展示的前端项目：用户在网页中手绘一条走势曲线，系统会在股票池里扫描最相似的 K 线窗口，并按相似度从高到低返回结果。

## 核心功能

- 手绘归一化曲线，支持预设形态快速起笔
- 对股票池做滑动窗口扫描，输出相似度、相关系数、形态贴合度、方向一致性
- 用专业 K 线图展示匹配结果，包含横纵坐标与鼠标悬浮信息
- 默认内置真实 A 股样本库，公开静态部署后可直接演示
- 预留 Tushare Token、代理 URL、股票代码与日期区间配置
- 结果区支持滚动展示，适合一次返回多条匹配结果

## 技术架构

- 前端：React 19 + TypeScript + Vite
- 图表：ECharts，按需在结果卡片中延迟加载
- 匹配逻辑：浏览器端执行，无需传统后端服务器
- 样本数据：构建期写入 `public/sample-market-snapshot.json`
- 可选实时数据：通过 Tushare API 获取；若公开部署在纯静态站点上，需要额外代理来规避浏览器跨域限制

## 为什么这样设计

Tushare 官方接口并不适合 GitHub Pages 这类纯静态网页直接在浏览器中调用，因此项目采用了双通道方案：

1. 演示通道：直接使用仓库内置的真实样本数据，站点上传后就能运行。
2. 扩展通道：保留 Tushare 配置入口；如果后续部署到支持无服务器函数的平台，或者自行提供代理 URL，就可以同步自定义股票池。

这意味着项目不需要自建传统后端服务器，但仍然兼顾了公开展示与后续扩展能力。

## 本地运行

```bash
npm install
npm run dev
```

构建生产版本：

```bash
npm run build
```

## 样本数据更新

项目提供了一个脚本从公开行情源生成演示用样本库：

```bash
python scripts/build_sample_market.py
```

输出文件：

- `public/sample-market-snapshot.json`

## GitHub 发布

仓库已经适配 GitHub Pages 自动部署。将整个项目上传到 GitHub 后：

1. 在仓库设置中进入 `Settings -> Pages`
2. 将 `Source` 设为 `GitHub Actions`
3. 推送代码到默认分支
4. 等待 Actions 中的 `Deploy static site to GitHub Pages` 工作流执行完成
5. GitHub 会自动生成公开访问地址

## 需要实时 Tushare 时

- 纯静态 GitHub Pages 适合演示样本模式
- 如果要让用户在公网环境下稳定使用自己的 Tushare Token，建议把同一套前端部署到 Vercel / Netlify，并额外提供一个无服务器代理接口

## 目录说明

- `src/App.tsx`：主界面与交互流程
- `src/components/PatternPad.tsx`：手绘曲线画板
- `src/components/KLineChart.tsx`：K 线结果图表
- `src/lib/matching.ts`：相似度匹配算法
- `src/lib/tushare.ts`：Tushare 数据接入封装
- `scripts/build_sample_market.py`：演示样本生成脚本
- `.github/workflows/deploy.yml`：GitHub Pages 自动部署工作流
""".strip() + "\n"


PROMPT_LINES = [
    "你是一名同时具备金融数据工程、前端产品设计、量化分析与 Web 部署经验的高级工程师。请为我从零实现一个“股票K线形态匹配工具”完整项目，并直接输出可运行代码、目录结构、测试方案和部署说明。项目必须面向金融学专业课程展示，整体风格要专业、美观、可信，不能做成通用 AI 模板页面。",
    "",
    "项目目标：",
    "做一个网页工具，让用户在页面上手绘一条曲线，系统自动在股票历史K线数据中搜索与该曲线走势最相似的股票片段，并按相似度从高到低展示结果。",
    "",
    "必须满足的功能要求：",
    "1. 股票价格数据支持通过 Tushare 获取，用户可以在页面内配置自己的 Tushare Token。",
    "2. 工具必须使用 Web UI 搭建，页面风格要美观、专业、偏金融终端 / 投研平台风格。",
    "3. 系统要考虑匹配速度和效率，能够尽可能准确匹配到与用户绘制曲线相似的股票走势，并显示相似度分数。",
    "4. 用户可以自主选择返回多少条匹配结果，结果需要按照相似度从高到低排列；如果页面展示不下，结果区域必须支持滚动。",
    "5. 展示出的K线图必须专业、美观，带横纵坐标；鼠标悬停时能够显示每日开盘价、收盘价、最高价、最低价、成交量等信息。",
    "6. 项目需要能够发布为公开网站，并且尽量不依赖传统后端服务器；如果 Tushare 的浏览器直连存在跨域限制，请你主动设计一个“静态前端主站 + 可选无服务器代理 / 可选离线样本库”的架构，确保项目既能公开展示，也能继续扩展。",
    "7. 页面中要明确区分“演示样本模式”和“实时数据模式”，避免公开部署后因为接口限制无法演示。",
    "8. Tushare Token、API Key 等敏感信息不能硬编码进公开仓库，必须只保存在本地浏览器或部署平台环境变量中。",
    "",
    "匹配算法要求：",
    "1. 用户绘制的曲线需要先进行归一化处理。",
    "2. 股票历史数据要采用滑动窗口方式逐段扫描。",
    "3. 相似度不能只看简单涨跌幅，至少结合以下因素：相关系数、形态距离、方向一致性。",
    "4. 要输出相似度总分，并尽量拆分出若干子指标，便于展示算法可信度。",
    "5. 在前端可承受的范围内优化性能，例如控制股票池、预处理窗口、延迟加载图表、避免不必要的重渲染。",
    "",
    "前端设计要求：",
    "1. 页面需要体现“金融研究工具”的气质，配色、排版、边框、数据卡片都要有专业感。",
    "2. 首页要有项目标题、简介、状态区域、参数区域、手绘区域、结果区域。",
    "3. 手绘区域支持鼠标拖拽绘制曲线，并提供若干预设形态按钮，例如 V 形反转、平台突破、回踩再起。",
    "4. 结果区域中的每个匹配结果都要包含：股票名称、股票代码、日期区间、相似度、子评分、K线图。",
    "5. 页面要兼容桌面端和移动端，至少保证移动端不崩溃、能滚动查看。",
    "",
    "技术实现建议：",
    "1. 优先采用 React + TypeScript + Vite。",
    "2. 图表优先采用 ECharts 或其他成熟金融图表方案。",
    "3. 默认内置一份真实或示例股票样本数据，保证项目在公开部署后无需配置任何私密信息也能立刻演示。",
    "4. 将 Tushare 接入设计成可切换的数据通道，而不是唯一依赖。",
    "5. 代码结构要清晰，至少拆分为：主页面、曲线绘制组件、K线图组件、匹配算法模块、数据获取模块、部署说明文档。",
    "",
    "最终交付要求：",
    "1. 输出完整项目代码。",
    "2. 输出 README，说明项目功能、架构、运行方式、部署方式。",
    "3. 输出测试方案，说明如何验证匹配功能、图表显示、滚动区域、Token 配置、公开部署可行性。",
    "4. 输出一份架构图说明，清楚展示前端、数据源、可选代理、样本数据之间的关系。",
    "5. 如果你判断“纯 GitHub Pages + 浏览器直连 Tushare”不可行，必须明确说明原因，并给出可落地替代方案，而不是忽略这个问题。",
]


FOLLOWUPS = [
    "请进一步优化匹配算法，说明你为什么选择“相关系数 + 形态误差 + 方向一致性”的组合评分。",
    "请补充一个适合课堂答辩的系统架构说明，重点解释为什么要采用“样本模式 + Tushare扩展模式”的双通道设计。",
    "请补充一份公开部署说明，分别说明 GitHub Pages 静态演示方案和 Vercel / Netlify 无服务器扩展方案。",
    "请给出 5 条可以用于课堂演示的手绘形态样例，并说明预期会匹配到什么样的K线走势。",
]


OPTIMIZATIONS = [
    "把功能需求、算法要求、UI要求、架构约束、交付清单拆开，减少 AI 漏项。",
    "提前点明 Tushare 的浏览器跨域风险，避免 AI 生成一个上线即失效的方案。",
    "强制加入演示样本模式，保证公开网站无需私密 Token 也能展示。",
    "强调结果排序、滚动展示、鼠标悬停明细、专业 K 线图等可被老师直接观察的细节。",
    "把安全性写进提示词，避免敏感 Token 被写死到仓库里。",
]


def write_readme(root: Path) -> None:
    readme_path = root / "README.md"
    readme_path.write_text(README_TEXT, encoding="utf-8")


def write_docx(root: Path) -> None:
    output = root / "docx" / "project_prompt_optimized.docx"
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(11)

    heading = doc.add_paragraph()
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = heading.add_run("K线图匹配工具项目提示词（优化版）")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.add_paragraph("用途：用于向 AI 提交高质量需求，使其一次性生成更完整、更可交付的课程项目。")
    doc.add_paragraph("适用场景：金融学专业小组作业、GitHub 展示、公开网站部署、课堂答辩演示。")

    doc.add_heading("一、推荐主提示词", level=1)
    for line in PROMPT_LINES:
      paragraph = doc.add_paragraph(line)
      if line.endswith("：") or line.endswith(":"):
        paragraph.runs[0].bold = True

    doc.add_heading("二、补充追问提示词", level=1)
    for item in FOLLOWUPS:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("三、这版提示词的优化重点", level=1)
    for item in OPTIMIZATIONS:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(output)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    write_readme(root)
    write_docx(root)


if __name__ == "__main__":
    main()
